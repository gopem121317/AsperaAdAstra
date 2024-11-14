import os
from docx.shared import Inches
from typing import List
import docx
from langchain_core.documents import Document
import docx2txt
import pandas as pd
from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from langchain_core.documents.base import Document as LDocument
import uuid
import re


def load_docx_unstructured(docx_file_path):
    doc = UnstructuredWordDocumentLoader(docx_file_path, mode="elements")
    data = doc.load()
    return data


class DocHandler:

    def __init__(self):
        # Initialize documents variables
        self.list_documents_parsed = []
        self.list_documents_docx = {}
        pass

    def split_word_document(self, word_document_path: str) -> List[dict]:
        """
        split word document to sub-documents based on their category

        :param word_document_path:
        :return:
        """
        doc_loader = UnstructuredWordDocumentLoader(word_document_path, mode='elements')
        docs_new = doc_loader.load()
        # generate unique id
        list_docs = list()
        table_count = 1
        for doc in docs_new:
            new_ele = {'id': str(uuid.uuid1()), 'document': doc}
            if doc.metadata['category'] == 'Table':
                doc.metadata['table_index'] = table_count
                table_count += 1
            list_docs.append(new_ele)

        # loading figure
        try:
            os.mkdir('./tmp/')
        except FileExistsError as e:
            Warning('tmp folder already exists')
        tmp_folder = f"./tmp/{str(uuid.uuid1())}"
        os.mkdir(tmp_folder)
        raw_text = docx2txt.process(word_document_path, tmp_folder)
        image_idx = 1
        for idx in range(len(list_docs)):
            if re.search(r'Figure [0-9]+:', list_docs[idx]['document'].page_content):
                image_path = f'{tmp_folder}/image{str(image_idx)}.jpeg'
                list_docs[idx]['image_path'] = image_path
                image_doc = self.load_image_simple(
                    caption=list_docs[idx]['document'].page_content,
                    image_path=word_document_path,
                    image_idx=image_idx,
                    document_name=os.path.basename(word_document_path)
                )
                list_docs.append({'id': uuid.uuid1(), 'document': image_doc})
                image_idx += 1

        # loading tables and save for later updating usage
        self.list_documents_docx[word_document_path] = docx.Document(word_document_path)
        self.list_documents_parsed.extend(list_docs)

    def load_image_simple(self, caption: str, image_path: str, image_idx: int, document_name: str = None):
        if document_name is None:
            document_name = os.path.basename(image_path)

        return Document(
            page_content=caption,
            metadata={"category": "Image",
                      "filename": document_name,
                      "source": image_path,
                      "image_index": image_idx}
        )


    def export_docs(self):
        """Export a list of parsed documents"""
        return self.list_documents_parsed

    def update_document(self, update_req, output_doc_path='./updated.docx'):
        """
        Given the update requirement, update corresponding document, and save to specified location.
        Each function call can only update one document but allows for multiple updates

        :param update_req: a list of update requirements saved in dictionary, each dictionary has an id and new document
        that is going to replace the existing one
        :param output_doc_path: save update docx to specified path
        :return: None
        """
        for update in update_req:
            old_ele_id = update['id']
            old_doc = self._look_up_ele_with_id(old_ele_id)
            # Update existing document
            target_file_name = old_doc['document'].metadata['source']
            doc_to_be_updated = self.list_documents_docx[target_file_name]

            if 'document' in update.keys():
                # where to be updated
                is_replaced = False
                for idx in range(len(doc_to_be_updated.paragraphs)):
                    if doc_to_be_updated.paragraphs[idx].text == old_doc['document'].page_content:
                        is_replaced = True
                        doc_to_be_updated.paragraphs[idx].text = update['document'].page_content
                        break
                if not is_replaced:
                    raise ValueError('Paragraph is not  found, fail to replace text paragraphs')

            elif 'image' in update.keys():
                # Update table
                is_replaced = False
                if old_doc['document'].metadata['category'] != 'Image':
                    raise TypeError('mismatch element category during update')
                target_image_index = old_doc['document'].metadata['image_index']
                image_count = 0
                for idx in range(len(doc_to_be_updated.paragraphs)):
                    if self._is_paragraph_image(doc_to_be_updated.paragraphs[idx]):
                        image_count += 1
                    if image_count == target_image_index:
                        # remove the existing image
                        for run in doc_to_be_updated.paragraphs[idx].runs:
                            doc_to_be_updated.paragraphs[idx]._p.remove(run.element)
                        # add the new image
                        doc_to_be_updated.paragraphs[idx].add_run().add_picture(update['image'], width=Inches(5.5))
                        break

            elif 'table' in update.keys():
                # update image
                table_index = old_doc['document'].metadata['table_index'] - 1
                old_table_obj = doc_to_be_updated.tables[table_index]
                if isinstance(update['table'], pd.DataFrame):
                    self._replace_table(old_table_obj, update['table'])
                else:
                    raise TypeError('Unexpected data type from update element')

            else:
                raise TypeError(f"Non-compatible new doc object")

        doc_to_be_updated.save(output_doc_path)

    def _look_up_ele_with_id(self, id):
        found_ele = None
        for element in self.list_documents_parsed:
            if id == element['id']:
                found_ele = element
                break
        if found_ele is None:
            raise KeyError('No id is found')
        return found_ele

    def _is_paragraph_image(self, paragraph):
        return 'Picture' in paragraph._element.xml

    def _replace_table(self, table_obj, new_table):
        """
        Replace a table in a Word document with data from an Excel file. The update is in-place to the input table
        object

        Parameters:
        word_file_path (str): Path to the Word document
        excel_file_path (str): Path to the Excel file
        sheet_name (str): Name of the Excel sheet containing the data
        table_index (int): Index of the table to replace (default is 0 for first table)
        """
        try:

            # Delete all rows except the first one (header)
            for _ in range(len(table_obj.rows) - 1):
                table_obj._element.remove(table_obj.rows[-1]._element)

            # Add new rows based on DataFrame content
            for _, row_data in new_table.iterrows():
                row_cells = table_obj.add_row().cells
                for col_idx, value in enumerate(row_data):
                    if col_idx < len(row_cells):  # Ensure we don't exceed the number of columns
                        row_cells[col_idx].text = str(value)

        except Exception as e:
            print(f"An error occurred: {str(e)}")
        return table_obj


if __name__ == "__main__":
    handler = DocHandler()
    handler.split_word_document(r'C:\Users\BNi\IdeaProjects\hackathon_AAA\DocUpdaterAAA\raw_data\model_doc\Time_Series_Model_Development_Report_refined.docx')
    df = pd.read_excel(r'C:\Users\BNi\IdeaProjects\hackathon_AAA\DocUpdaterAAA\raw_data\model_doc\Generated_Analysis_Tables.xlsx', sheet_name='Dataset Summary')
    update_req = [{'id': handler.list_documents_parsed[0]['id'], 'document': handler.list_documents_parsed[1]['document']},
                  {'id': handler.list_documents_parsed[34]['id'], 'table': df},
                  {'id': handler.list_documents_parsed[-1]['id'], 'image': r'C:\Users\BNi\IdeaProjects\hackathon_AAA\DocUpdaterAAA\raw_data\model_doc\Business X_scatter.jpeg'}]
    handler.update_document(update_req)
    handler.export_docs()