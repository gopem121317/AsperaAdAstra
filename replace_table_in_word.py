from docx import Document
import pandas as pd
import os

def replace_table_from_excel(word_file_path, excel_file_path, sheet_name, table_index=0):
    """
    Replace a table in a Word document with data from an Excel file.
    
    Parameters:
    word_file_path (str): Path to the Word document
    excel_file_path (str): Path to the Excel file
    sheet_name (str): Name of the Excel sheet containing the data
    table_index (int): Index of the table to replace (default is 0 for first table)
    """
    try:
        # Read the Excel file
        df = pd.read_excel(excel_file_path, sheet_name=sheet_name)
        
        # Load the Word document
        doc = Document(word_file_path)
        
        # Check if the document has tables
        if len(doc.tables) <= table_index:
            raise IndexError(f"Table index {table_index} is out of range. Document has {len(doc.tables)} tables.")
        
        # Get the target table
        table = doc.tables[table_index]
        
        # Delete all rows except the first one (header)
        for _ in range(len(table.rows) - 1):
            table._element.remove(table.rows[-1]._element)
        
        # Add new rows based on DataFrame content
        for _, row_data in df.iterrows():
            row_cells = table.add_row().cells
            for col_idx, value in enumerate(row_data):
                if col_idx < len(row_cells):  # Ensure we don't exceed the number of columns
                    row_cells[col_idx].text = str(value)
        
        # Create output filename
        file_name, file_ext = os.path.splitext(word_file_path)
        output_path = f"{file_name}_updated{file_ext}"
        
        # Save the modified document
        doc.save(output_path)
        print(f"Document successfully updated and saved as: {output_path}")
        
    except Exception as e:
        print(f"An error occurred: {str(e)}")

# Example usage
if __name__ == "__main__":
    # Example paths - replace with your actual file paths
    word_doc = "input.docx"
    excel_file = "data.xlsx"
    sheet = "Sheet1"
    
    replace_table_from_excel(
        word_file_path=word_doc,
        excel_file_path=excel_file,
        sheet_name=sheet,
        table_index=0  # Replace first table
    )